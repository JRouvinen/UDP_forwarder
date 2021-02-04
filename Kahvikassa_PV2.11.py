#-------------------------------------
#Dev notes
#Main program for Kahvikerho aplication, driven from ttk_gui_v0823_dev branch
#-------------------------------------
#List of imports:
import sys
try:
    import ast
    import io
    import logging
    import os
    import random
    import shutil
    import sqlite3
    import subprocess
    import tkinter as tk
    import tkinter.font as tkFont
    import math
    from random import randint
    from time import sleep
    # importing strftime function to
    # retrieve system's time
    from time import strftime
    from tkinter import *
    from tkinter import filedialog
    from tkinter import messagebox
    from tkinter import ttk
    from tkinter.ttk import *

    from PIL import Image, ImageTk
    from docx import Document
    from ttkthemes import ThemedTk
except ImportError as e:
    print(e)
    pass
#---------------------Version notes and important things------------------------------------
#TODO LIST:
#
#
#
#
#
#LIST OF BUGS:
#Serial
#1: When updating namelist program restart doesn't work -> check that restart actually happens!
#2: Check is second info text bar needed in main window
#3: After month change payments won't work anymore, message: "*savepayment* ERROR. Reason: no such table: debts"
#4: If debts table has payment 5.8888888 system won't round up values.

#---------------------------------------------------------------------------------------------

#Version notes
#Version 0.3 has SQLite support
#Version 0.4 has docx support
#version 0.5 created decent class and def __init__ to support better startup procedures
#version 0.6 has software key checking written into it
#version 0.61 has setting window impedded and local DB files are read from \data folder
#version 0.62 has name_management impedded into main code
#version 0.63 has loggin support impedded
#version 0.64 has checks and questions for key check and for giving key -->developt to a launcher?
#version 0.70 has own GUI to check and give program key.
#version 0.71 has exeptions and further logging written into it
#version 0.8 has debt management impedded into the main code
#version 0.81 has improved devaddpayments
#version 0.82 has impedded history tool and revorked settings and devtools windows
#version 0.821 has better settings window and new Devtools menu on main gui topbar
#version 0.822 has better compatibily with lower resolution machines
#version 0.823 has imporoved scaling of the payment buttons and added seurity features
#version 0.823FIN has Finnis language as default
#version 0.824FIN has better debug tools implemented and some code cleaning started
#version 0.825FIN has more cleaning done and updated debt management features
#version 0.8230FIN_RC has cleaned redundant code and hidden debug features
#version 0.8231FIN_RC has cleaned main GUI, added two new payment buttons and history db is saved now by date+month
#version 0.8232FIN_RC has history db month check and db update. Name input will strip excess white spaces out from name.
# Cleaned payment editing flow -> each payment is processed on save.
#version 0.8232AFIN_RC when error occurs program shows messagebox and creates error file to desktop
#version 0.8233AFIN_RC has updated debts history process -> saves single datetime instead of date and time,
# also erasing all debts saves info to debt history db
#version 0.8233AFIN_RC branch has made to official delivery named: Kahvikassa_PV2
#version Kahvikassa_PV2.1 has following bug fixes: 3
#version Kahvikassa_PV2.11 has following bug fixes: 4

#---------------------Define GUI settings------------------------------------

class mainprogram(object):

    def launchergui(self):
    #-----------START OF - Key check-------------------------------
        print('Checking existing keys from DB')
        logging.info('Launcher started')
        conn = sqlite3.connect('DB\people.db')
        c = conn.cursor()
        logging.debug('Connected to DB people.db')
        c.execute('''CREATE TABLE IF NOT EXISTS softwarekey
                         (date, key, machineID)''')
        #Get the software key from DB
        c.execute("SELECT key FROM softwarekey")

        get_key = c.fetchall()
        softw_key = str(get_key)
        softw_key = softw_key.replace('[','')
        softw_key = softw_key.replace(']','')
        softw_key = softw_key.replace('\'','')
        softw_key = softw_key.replace(',','')
        softw_key = softw_key.replace('(','')
        softw_key = softw_key.replace(')','')
        logging.debug('Key from DB:')
        logging.debug(softw_key)

        #Get the machine ID from DB
        c.execute("SELECT machineID FROM softwarekey")

        get_saved_mID = c.fetchall()
        saved_mID = str(get_saved_mID)
        saved_mID = saved_mID.replace('[','')
        saved_mID = saved_mID.replace(']','')
        saved_mID = saved_mID.replace('\'','')
        saved_mID = saved_mID.replace(',','')
        saved_mID = saved_mID.replace('(','')
        saved_mID = saved_mID.replace(')','')

        conn.commit()
        #close the connection
        conn.close()
        #end of key check
        logging.debug('Closed connection to DB people.db')
        key = softw_key
        self.keycheck_var = False
        #define used key library in here
        self.bravolib = {
            "Created": 'ACS KEY GENERATOR',
            "Version": ' V3-A10',
            "Serial": '17112020-HFO5L-C5A86FA6',
            'A': 'Y',
            'B': 'W',
            'C': 'R',
            'D': '8',
            'E': 'D',
            'F': '9',
            'G': 'N',
            'H': 'I',
            'I': 'A',
            'J': 'E',
            'K': 'P',
            'L': 'S',
            'M': '4',
            'N': '0',
            'O': 'K',
            'P': 'O',
            'Q': '1',
            'R': 'U',
            'S': 'Z',
            'T': 'M',
            'U': 'B',
            'V': '3',
            'W': 'F',
            'X': 'Q',
            'Y': 'J',
            'Z': 'L',
            '1': '7',
            '2': 'X',
            '3': '6',
            '4': 'V',
            '5': 'H',
            '6': 'T',
            '7': '5',
            '8': '2',
            '9': 'C',
            '0': 'G',
            'A1': 16,
            'A2': 18,
            'A3': 10,
            'A4': 4,
            'A5': 19,
            'B1': 19,
            'B2': 1,
            'B3': 12,
            'B4': 5,
            'B5': 3,
            'C1': 18,
            'C2': 19,
            'C3': 10,
            'C4': 6,
            'C5': 0,
            'D1': 29,
            'D2': 27,
            'D3': 5,
            'D4': 17,
            'D5': 10
        }

        if key == '':
            print('No keys found! Starting key inquiry.')
            self.keycheck_var = False
            logging.info('No key found!')
            logging.info('Starting key request')
            #launcher_gui = ThemedTk(theme="arc")
            self.launcher_gui = Toplevel(gui)

            #Set GUI parameters - Title
            self.launcher_gui.title("ACS Program Launcher")
            #Set GUI parameters - Window size x+y
            self.launcher_gui.geometry("450x320")
            #frame
            kg_fr1=ttk.Frame(self.launcher_gui,relief="raised", borderwidth=1)
            kg_fr1.pack(expand=0, fill="both")
            kg_fr2=ttk.Frame(self.launcher_gui,relief="raised", borderwidth=1)
            kg_fr2.pack(expand=0, fill="both")
            #style
            launcherstyle = ttk.Style()
            launcherstyle.configure("Linfo.TLabel", font=("calibri", 20, "bold"))
            launcherstyle.configure("Ltxt.TLabel", font=("calibri", 15))
            #Variables
            self.linfotxt = StringVar()
            self.linfotxt.set("No key found!")
            #Start the GUI/Create main loop
            l_ilbl1 = ttk.Label(kg_fr1, textvariable=self.linfotxt, style="Linfo.TLabel").pack()
            l_lbl1 = ttk.Label(kg_fr1, text="To obtain key report serial below to the developer.", style="Ltxt.TLabel").pack()
            l_lbl2 = ttk.Label(kg_fr1, text=self.curr_mID, style="Ltxt.TLabel").pack()
            l_lbl3 = ttk.Label(kg_fr1, text="E-mail:", style="Ltxt.TLabel").pack()
            l_lbl4 = ttk.Label(kg_fr1, text="(jmrouvinen@gmail.com)",style="Ltxt.TLabel").pack()
            l_lbl5 = ttk.Label(kg_fr1, text="",style="Ltxt.TLabel").pack()
            l_lbl6 = ttk.Label(kg_fr1, text="Key:",style="Ltxt.TLabel").pack()
            l_ent1 = self.launcher_ent = ttk.Entry(kg_fr1, width=25)
            self.launcher_ent.pack()
            l_lbl7 = ttk.Label(kg_fr1, text="",style="Ltxt.TLabel").pack()
            l_btn1 = ttk.Button(kg_fr1, text="OK", command=self.keycheck).pack()
            l_btn2 = ttk.Button(kg_fr1, text="Exit",command=self.launcher_gui.quit).pack()
            l_lbl8 = ttk.Label(kg_fr1, text="",style="Ltxt.TLabel").pack()

        else:
            print('Key found from DB, verifying.')
            logging.info('Key found from DB, verifying key')
            logging.debug('Starting launcher')
            #launcher_gui = ThemedTk(theme="arc")
            self.launcher_gui = Toplevel(gui)
            #Set GUI parameters - Title
            self.launcher_gui.title("ACS Program Launcher")
            #Set GUI parameters - Window size x+y
            self.launcher_gui.geometry("450x200")

            #frame
            kg_fr1=ttk.Frame(self.launcher_gui,relief="raised", borderwidth=1)
            kg_fr1.pack(expand=0, fill="both")

            #style
            launcherstyle = ttk.Style()
            launcherstyle.configure("Linfo.TLabel", font=("calibri", 20, "bold"))
            launcherstyle.configure("Ltxt.TLabel", font=("calibri", 15))
            #Variables
            self.linfotxt = StringVar()
            self.linfotxt.set("Verifying key...")
            self.linfotxt2 = StringVar()
            self.linfotxt2.set("")
            #load image
            load = Image.open("data\ACS_logo_small_75pxl.png")
            render = ImageTk.PhotoImage(load)
            img = Label(kg_fr1, image=render)
            img.image = render
            img.place(x=0, y=0)
            #Labels
            l_lbl1 = ttk.Label(kg_fr1, text="Key found.", style="Linfo.TLabel").pack()
            l_ilbl1 = ttk.Label(kg_fr1, textvariable=self.linfotxt, style="Ltxt.TLabel").pack()
            l_ilbl2 = ttk.Label(kg_fr1, textvariable=self.linfotxt2, style="Ltxt.TLabel").pack()
            l_lbl7 = ttk.Label(kg_fr1, text="",style="Ltxt.TLabel").pack()
            l_btn1 = ttk.Button(kg_fr1, text="Start", command=self.startprogram).pack()
            l_btn2 = ttk.Button(kg_fr1, text="Exit", command=self.launcher_gui.quit).pack()
            l_lbl7 = ttk.Label(kg_fr1, text="",style="Ltxt.TLabel").pack()
            #UUID check starts here
            logging.debug('Prosessing UUID')

            self.pros_UUID = self.curr_mID.replace('-','')

            self.A_seq1_1 = str(self.pros_UUID[self.bravolib['A1']])
            self.A_seq1_2 = str(self.pros_UUID[self.bravolib['A2']])
            self.A_seq1_3 = str(self.pros_UUID[self.bravolib['A3']])
            self.A_seq1_4 = str(self.pros_UUID[self.bravolib['A4']])
            self.A_seq1_5 = str(self.pros_UUID[self.bravolib['A5']])
            self.UUID_seq1 = self.A_seq1_1+self.A_seq1_2+self.A_seq1_3+self.A_seq1_4+self.A_seq1_5
            #print(self.UUID_seq1)

            self.B_seq1_1 = str(self.pros_UUID[self.bravolib['B1']])
            self.B_seq1_2 = str(self.pros_UUID[self.bravolib['B2']])
            self.B_seq1_3 = str(self.pros_UUID[self.bravolib['B3']])
            self.B_seq1_4 = str(self.pros_UUID[self.bravolib['B4']])
            self.B_seq1_5 = str(self.pros_UUID[self.bravolib['B5']])
            self.UUID_seq2 = self.B_seq1_1+self.B_seq1_2+self.B_seq1_3+self.B_seq1_4+self.B_seq1_5
            #print(self.UUID_seq2)

            self.C_seq1_1 = str(self.pros_UUID[self.bravolib['C1']])
            self.C_seq1_2 = str(self.pros_UUID[self.bravolib['C2']])
            self.C_seq1_3 = str(self.pros_UUID[self.bravolib['C3']])
            self.C_seq1_4 = str(self.pros_UUID[self.bravolib['C4']])
            self.C_seq1_5 = str(self.pros_UUID[self.bravolib['C5']])
            self.UUID_seq3 = self.C_seq1_1+self.C_seq1_2+self.C_seq1_3+self.C_seq1_4+self.C_seq1_5
            #print(self.UUID_seq3)

            self.D_seq1_1 = str(self.pros_UUID[self.bravolib['D1']])
            self.D_seq1_2 = str(self.pros_UUID[self.bravolib['D2']])
            self.D_seq1_3 = str(self.pros_UUID[self.bravolib['D3']])
            self.D_seq1_4 = str(self.pros_UUID[self.bravolib['D4']])
            self.D_seq1_5 = str(self.pros_UUID[self.bravolib['D5']])
            self.UUID_seq4 = self.D_seq1_1+self.D_seq1_2+self.D_seq1_3+self.D_seq1_4+self.D_seq1_5

            self.Full_key_seq = self.UUID_seq1+'-'+self.UUID_seq2+'-'+self.UUID_seq3+'-'+self.UUID_seq4

            logging.debug('Machine UUID key created')
            logging.debug(self.Full_key_seq)

            if self.Full_key_seq != softw_key:
                logging.debug('NOT CORRECT KEY!')
                self.linfotxt.set("NOT CORRECT KEY!")
                print('NOT CORRECT KEY! Stopping program...')
                #Deleting current invalid key from DB
                key_to_save = ''
                date = strftime("%d-%m-%Y")
                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()
                logging.debug('Connected to DB people.db')
                sqlite_delete_key = """DELETE FROM softwarekey;"""
                c.execute(sqlite_delete_key)
                logging.info('Key removed from DB')
                conn.commit()
                #close the connection
                conn.close()
                self.program()
                sleep(10)
                sys.exit()
            else:
                logging.debug('UUID vs KEY MATCH!')
                print('Correct key, launching GUI...')
                self.linfotxt.set("CORRECT KEY \n Press 'START' to start program.")
                self.startprogram()

    def startprogram(self):

        logging.info('Program launched.')
        self.program()
        self.launcher_gui.destroy()

    def keycheck(self):
           # this part needs to be separated into its own function!
            launcher_input = str(self.launcher_ent.get())
            if launcher_input == '':
                self.linfotxt.set("GIVE KEY!")
            else:
                ui = launcher_input
                #UUID check starts here
                logging.debug('Prosessing UUID')

                self.pros_UUID = self.curr_mID.replace('-','')

                self.A_seq1_1 = str(self.pros_UUID[self.bravolib['A1']])
                self.A_seq1_2 = str(self.pros_UUID[self.bravolib['A2']])
                self.A_seq1_3 = str(self.pros_UUID[self.bravolib['A3']])
                self.A_seq1_4 = str(self.pros_UUID[self.bravolib['A4']])
                self.A_seq1_5 = str(self.pros_UUID[self.bravolib['A5']])
                self.UUID_seq1 = self.A_seq1_1+self.A_seq1_2+self.A_seq1_3+self.A_seq1_4+self.A_seq1_5
                #print(self.UUID_seq1)

                self.B_seq1_1 = str(self.pros_UUID[self.bravolib['B1']])
                self.B_seq1_2 = str(self.pros_UUID[self.bravolib['B2']])
                self.B_seq1_3 = str(self.pros_UUID[self.bravolib['B3']])
                self.B_seq1_4 = str(self.pros_UUID[self.bravolib['B4']])
                self.B_seq1_5 = str(self.pros_UUID[self.bravolib['B5']])
                self.UUID_seq2 = self.B_seq1_1+self.B_seq1_2+self.B_seq1_3+self.B_seq1_4+self.B_seq1_5
                #print(self.UUID_seq2)

                self.C_seq1_1 = str(self.pros_UUID[self.bravolib['C1']])
                self.C_seq1_2 = str(self.pros_UUID[self.bravolib['C2']])
                self.C_seq1_3 = str(self.pros_UUID[self.bravolib['C3']])
                self.C_seq1_4 = str(self.pros_UUID[self.bravolib['C4']])
                self.C_seq1_5 = str(self.pros_UUID[self.bravolib['C5']])
                self.UUID_seq3 = self.C_seq1_1+self.C_seq1_2+self.C_seq1_3+self.C_seq1_4+self.C_seq1_5
                #print(self.UUID_seq3)

                self.D_seq1_1 = str(self.pros_UUID[self.bravolib['D1']])
                self.D_seq1_2 = str(self.pros_UUID[self.bravolib['D2']])
                self.D_seq1_3 = str(self.pros_UUID[self.bravolib['D3']])
                self.D_seq1_4 = str(self.pros_UUID[self.bravolib['D4']])
                self.D_seq1_5 = str(self.pros_UUID[self.bravolib['D5']])
                self.UUID_seq4 = self.D_seq1_1+self.D_seq1_2+self.D_seq1_3+self.D_seq1_4+self.D_seq1_5

                self.Full_key_seq = self.UUID_seq1+'-'+self.UUID_seq2+'-'+self.UUID_seq3+'-'+self.UUID_seq4


                if ui != self.Full_key_seq:
                    self.linfotxt.set("NOT CORRECT KEY!")
                    logging.info('No valid key')


                else:
                    self.linfotxt.set("Validating key")
                    logging.info('Valid key')
                    sleep(1)
                    self.linfotxt.set("Starting program...")
                    logging.info('Starting program...')
                    sleep(2)
                    key_to_save = ui
                    #print(key_to_save)
                    date = strftime("%d-%m-%Y")
                    conn = sqlite3.connect('DB\people.db')
                    c = conn.cursor()
                    logging.debug('Connected to DB people.db')
                    #print("Connected to SQLite DB")
                    sqlite_insert_with_param = """INSERT INTO softwarekey
                      (date, key)VALUES (?, ?);"""
                    data_tuple = (date, key_to_save)

                    #insert row
                    c.execute(sqlite_insert_with_param, data_tuple)
                    logging.info('Key saved to DB')
                    logging.debug(key_to_save)
                    conn.commit()
                    #close the connection
                    conn.close()
                    logging.debug('Closed connection to DB people.db')
                    self.program()
                    self.launcher_gui.destroy()
    #-----------END OF - Key check-------------------------------

    def program(self):
        logging.info(ver)
    #-----------START OF - GUI definitions, frames etc.-------------------------------

        #Define frames for GUI
        fr1=ttk.Frame(gui,relief="flat", borderwidth=1)
        fr1.pack(expand=0, fill="both")

        fr2=ttk.Frame(gui,relief="flat", borderwidth=1)
        fr2.pack(expand=0, fill="both")

        fr3=ttk.Frame(gui,relief="flat", borderwidth=0)
        fr3.pack(expand=0, fill="both")

        fr4=ttk.Frame(gui,relief="flat", borderwidth=0)
        fr4.pack(expand=0, fill="both", anchor="center")

        fr5=ttk.Frame(gui,relief="flat", borderwidth=0)
        fr5.pack(expand=0, fill="both")

        fr6=ttk.Frame(gui,relief="flat", borderwidth=0)
        fr6.pack(expand=0, fill="both", anchor="center")

#------------------Define variables-------------------------------

        self.debt_to_add = 0.0
        self.getname = ""
        self.version = ver
#------------------Menubar butons-------------------------------
        logging.info('Starting main GUI')
        logging.info('Program version: '+self.version)
        menubar = Menu(gui)
        filemenu = Menu(menubar, tearoff=0)
        filemenu.add_command(label="Hallitse velkoja", command=self.debtmang_pswrd)
        filemenu.add_command(label="Hallitse nimilistaa", command=self.nm_mngmt_pswrd)
        #filemenu.add_command(label="Manage Names", command=self.nm_mngmt)
        filemenu.add_command(label="Tulosta velkalista", command=self.print)
        filemenu.add_command(label="Vie velkalista Exceliin", command=self.exporttoexl)
        filemenu.add_command(label="Tarkastele historia tietoja", command=self.paymenthistory)

        filemenu.add_separator()

        filemenu.add_command(label="Exit", command=gui.quit)
        menubar.add_cascade(label="File", menu=filemenu)


        helpmenu = Menu(menubar, tearoff=0)
        helpmenu.add_command(label="Help Index", command=self.helpindx)
        helpmenu.add_command(label="About...", command=self.about)
        menubar.add_cascade(label="Help", menu=helpmenu)

        gui.config(menu=menubar)
#---------------------Define styling options ----------------------------------
        #Main window styles
        style = ttk.Style()
        style.configure("Main.TLabel", font=("calibri", 20, "bold"),fill='BOTH')
        style.configure("Clock.TLabel", font=("calibri", 15, "bold"))
        style.configure("Sec.TLabel", font=("calibri", 15, "bold"))
        style.configure("Empty.TLabel", font=("calibri", 8),fill='BOTH')
        style.configure("SuperEmpty.TLabel", font=("calibri", 58),fill='BOTH')
        style.configure("EmptyBTN.TLabel", font=("calibri", 12, "bold"))
        style.configure("Info.TLabel", font=("calibri", 12, "bold"))
        style.configure("Other.TLabel", font=("calibri", 10, "bold"))
        style.configure("Browse.TLabel", font=("calibri", 8, "bold"), fill='BOTH')

        style.configure("Basic.TButton", font=("calibri", 12, "bold"),fill='BOTH')
        style.configure("Small.TButton", font=("calibri", 10, "bold"),fill='BOTH')

        #settings window styles

#---------------------Define Top label on window (fr1) ------------------------------
        #Top label for window

        ttk.Label(fr1, text="IOK Kahvikerhon Kahvikassa", style="Main.TLabel").pack()
        load = Image.open("data\mainw_logo.jpg")
        render = ImageTk.PhotoImage(load)
        img = Label(fr1, image=render)
        img.image = render
        img.place(x=0, y=0)
#-----------START OF - GUI definitions, frames etc.-------------------------------

#---------------------Define date and clock------------------------------------
    #Date
        def date():
            string = strftime("%d-%m-%Y")
            lbl1.config(text = string)
            lbl1.after(1000, date)

        lbl1 = ttk.Label(fr1, style="Clock.TLabel")

        # Placing clock at the centre
        # of the tk window
        lbl1.pack(anchor = 'center')
        date()


        #Clock
        def time():
            global clk_style
            string = strftime('%H:%M:%S')
            lbl.config(text = string)
            lbl.after(1000, time)

        lbl = ttk.Label(fr1, style="Clock.TLabel")

        # Placing clock at the centre
        # of the tk window
        lbl.pack(anchor = 'center')
        time()
        ttk.Label(fr1, text="Info", style="Sec.TLabel").pack()
        ttk.Label(fr1, textvariable=self.infotext,style="Info.TLabel").pack()
        ttk.Label(fr1, textvariable=self.infotext2, style="Info.TLabel").pack()
        ttk.Label(fr1, text="", style="Sec.TLabel").pack()
    #---------------------Define Combobox (fr2)------------------------------------

        #Define Label and combobox
        ttk.Label(fr3, text="Syötä nimesi", style="Sec.TLabel").pack()
        #Define Combobox

        class AutocompleteEntry(tk.Entry):
                """
                Subclass of tk.Entry that features autocompletion.

                To enable autocompletion use set_completion_list(list) to define
                a list of possible strings to hit.
                To cycle through hits use down and up arrow keys.
                """
                def set_completion_list(self, completion_list):
                        self._completion_list = sorted(completion_list, key=str.lower) # Work with a sorted list
                        self._hits = []
                        self._hit_index = 0
                        self.position = 0
                        self.bind('<KeyRelease>', self.handle_keyrelease)

                def autocomplete(self, delta=0):
                        """autocomplete the Entry, delta may be 0/1/-1 to cycle through possible hits"""
                        if delta: # need to delete selection otherwise we would fix the current position
                                self.delete(self.position, tk.END)
                        else: # set position to end so selection starts where textentry ended
                                self.position = len(self.get())
                        # collect hits
                        _hits = []
                        for element in self._completion_list:
                                if element.lower().startswith(self.get().lower()):  # Match case-insensitively
                                        _hits.append(element)
                        # if we have a new hit list, keep this in mind
                        if _hits != self._hits:
                                self._hit_index = 0
                                self._hits=_hits
                        # only allow cycling if we are in a known hit list
                        if _hits == self._hits and self._hits:
                                self._hit_index = (self._hit_index + delta) % len(self._hits)
                        # now finally perform the auto completion
                        if self._hits:
                                self.delete(0,tk.END)
                                self.insert(0,self._hits[self._hit_index])
                                self.select_range(self.position,tk.END)

                def handle_keyrelease(self, event):
                        """event handler for the keyrelease event on this widget"""
                        if event.keysym == "BackSpace":
                                self.delete(self.index(tk.INSERT), tk.END)
                                self.position = self.index(tk.END)
                        if event.keysym == "Left":
                                if self.position < self.index(tk.END): # delete the selection
                                        self.delete(self.position, tk.END)
                                else:
                                        self.position = self.position-1 # delete one character
                                        self.delete(self.position, tk.END)
                        if event.keysym == "Right":
                                self.position = self.index(tk.END) # go to end (no selection)
                        if event.keysym == "Down":
                                self.autocomplete(1) # cycle to next hit
                        if event.keysym == "Up":
                                self.autocomplete(-1) # cycle to previous hit
                        if len(event.keysym) == 1 or event.keysym in tk_umlauts:
                                self.autocomplete()

        class AutocompleteCombobox(ttk.Combobox):

                def set_completion_list(self, completion_list):
                        """Use our completion list as our drop down selection menu, arrows move through menu."""
                        self._completion_list = sorted(completion_list, key=str.lower) # Work with a sorted list
                        self._hits = []
                        self._hit_index = 0
                        self.position = 0
                        self.bind('<KeyRelease>', self.handle_keyrelease)
                        self['values'] = self._completion_list  # Setup our popup menu

                def autocomplete(self, delta=0):
                        """autocomplete the Combobox, delta may be 0/1/-1 to cycle through possible hits"""
                        if delta: # need to delete selection otherwise we would fix the current position
                                self.delete(self.position, tk.END)
                        else: # set position to end so selection starts where textentry ended
                                self.position = len(self.get())
                        # collect hits
                        _hits = []
                        for element in self._completion_list:
                                if element.lower().startswith(self.get().lower()): # Match case insensitively
                                        _hits.append(element)
                        # if we have a new hit list, keep this in mind
                        if _hits != self._hits:
                                self._hit_index = 0
                                self._hits=_hits
                        # only allow cycling if we are in a known hit list
                        if _hits == self._hits and self._hits:
                                self._hit_index = (self._hit_index + delta) % len(self._hits)
                        # now finally perform the auto completion
                        if self._hits:
                                self.delete(0,tk.END)
                                self.insert(0,self._hits[self._hit_index])
                                self.select_range(self.position,tk.END)

                def handle_keyrelease(self, event):
                        """event handler for the keyrelease event on this widget"""
                        if event.keysym == "BackSpace":
                                self.delete(self.index(tk.INSERT), tk.END)
                                self.position = self.index(tk.END)
                        if event.keysym == "Left":
                                if self.position < self.index(tk.END): # delete the selection
                                        self.delete(self.position, tk.END)
                                else:
                                        self.position = self.position-1 # delete one character
                                        self.delete(self.position, tk.END)
                        if event.keysym == "Right":
                                self.position = self.index(tk.END) # go to end (no selection)
                        if len(event.keysym) == 1:
                                self.autocomplete()
                        # No need for up/down, we'll jump to the popup
                        # list at the position of the autocompletion

        def test(test_list):
                """Run a mini application to test the AutocompleteEntry Widget."""
                #global self.getname
                #root = tk.Tk(className=' AutocompleteEntry demo')
                #entry = AutocompleteEntry(topframe)
                #entry.set_completion_list(test_list)
                #entry.pack()
                #entry.focus_set()
                combo = AutocompleteCombobox(fr3, justify = 'center', font=("calibri", 20, "bold"))
                combo.set_completion_list(test_list)
                combo.pack()
                combo.focus_set()
                self.getname = combo
                # I used a tiling WM with no controls, added a shortcut to quit
                #root.bind('<Control-Q>', lambda event=None: root.destroy())
                #root.bind('<Control-q>', lambda event=None: root.destroy())
                #root.mainloop()

        if __name__ == '__main__':

#-----------CREATE PAYMENT HISTORY DB-------------------------------------
            date_my = strftime("%m-%Y")
            self.debt_history = date_my+"_debts.db"
            #print(self.debt_history)
            conn = sqlite3.connect('history\\'+self.debt_history)
            logging.info('Connecting to DB debts:')
            logging.info('history\\'+self.debt_history)
            c = conn.cursor()
            #create table 'people'
            c.execute('''CREATE TABLE IF NOT EXISTS debts
                         (datetime, name, debt, payment)''')

            conn.commit()
            #close the connection
            conn.close()
#-----------CREATE PEOPLE DB-------------------------------------
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            logging.debug('Connecting to DB people')
            #create table 'people'
            c.execute('''CREATE TABLE IF NOT EXISTS people
                         (date, time, name, debt, payment)''')

            #update payment section to 0
            c.execute('UPDATE people SET payment=0')

            c.execute("SELECT name FROM people")
            rows = c.fetchall()
            #print(rows)
            name_list = []
            for row in rows:
                name_list += row

            #print(name_list)
            #commit the changes to db
            conn.commit()
            #close the connection
            conn.close()
            logging.debug('Closed DB people')
            self.test_list = name_list
            test(self.test_list)

            #These lines are not needed for SQLite support, only if you want to use plain text file for names
            #test_list = open("people.txt", "r")
            #test(test_list)
            #test_list.close()

#-----------Cleaning DB for spurious debt values-----------------------------------
            logging.debug('Cleaning up DB people - START')
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            c.execute("SELECT round (debt,1) FROM people")

            rows = c.fetchall()

            for row in rows:
                c.execute("UPDATE people SET debt=round(debt,1)")
                #print(row)
                conn.commit()
            conn.close()
            logging.debug('Cleaning up DB people - DONE')
#------------------------------------------------------------------------------------

            ttk.Label(fr3, text="", style="Empty.TLabel").pack()
            dbt = ttk.Button(fr3, text="Tarkista velan määrä", style="Basic.TButton", command=self.debtquery).pack()
            ttk.Label(fr3, text="", style="Empty.TLabel").pack()
            #-------------Button layout section starts here (fr3+4)----------------------
            #Define Label and combobox
            ttk.Label(fr4, text="Syötä maksu", style="Sec.TLabel").pack()
            #ttk.Label(fr4, text="", style="Empty.TLabel").pack()
            #Define buttons
            ttk.Label(fr5, text="", style="SuperEmpty.TLabel").pack()
            pbt1 = ttk.Button(fr5, text="0,2 €", style="Basic.TButton", command=self.payment02).place(relx="0.2", rely="0.2", anchor="center")
            pbt2 = ttk.Button(fr5, text="0,4 €",  style="Basic.TButton", command=self.payment04).place(relx="0.3", rely="0.2", anchor="center")
            pbt3 = ttk.Button(fr5, text="0,5 €",  style="Basic.TButton", command=self.payment05).place(relx="0.4", rely="0.2", anchor="center")
            pbt4 = ttk.Button(fr5, text="0,6 €",  style="Basic.TButton", command=self.payment06).place(relx="0.5", rely="0.2", anchor="center")
            pbt5 = ttk.Button(fr5, text="0,7 €",  style="Basic.TButton", command=self.payment07).place(relx="0.6", rely="0.2", anchor="center")
            pbt6 = ttk.Button(fr5, text="0,8 €",  style="Basic.TButton", command=self.payment08).place(relx="0.7", rely="0.2", anchor="center")
            pbt7 = ttk.Button(fr5, text="1,0 €",  style="Basic.TButton", command=self.payment10).place(relx="0.8", rely="0.2", anchor="center")

            pbt8 = ttk.Button(fr5, text="1,2 €",  style="Basic.TButton", command=self.payment12).place(relx="0.2", rely="0.5", anchor="center")
            pbt9 = ttk.Button(fr5, text="1,4 €",  style="Basic.TButton", command=self.payment14).place(relx="0.3", rely="0.5", anchor="center")
            pbt10 = ttk.Button(fr5, text="1,6 €",  style="Basic.TButton", command=self.payment16).place(relx="0.4", rely="0.5", anchor="center")
            pbt11 = ttk.Button(fr5, text="1,8 €",  style="Basic.TButton", command=self.payment18).place(relx="0.5", rely="0.5", anchor="center")
            pbt12 = ttk.Button(fr5, text="2,0 €",  style="Basic.TButton", command=self.payment20).place(relx="0.6", rely="0.5", anchor="center")
            pbt13 = ttk.Button(fr5, text="2,2€",  style="Basic.TButton", command=self.payment22).place(relx="0.7", rely="0.5", anchor="center")
            pbt14 = ttk.Button(fr5, text="2,4 €",  style="Basic.TButton", command=self.payment24).place(relx="0.8", rely="0.5", anchor="center")
            pbt15 = ttk.Button(fr5, text="2,6 €", style="Basic.TButton", command=self.payment26).place(relx="0.2", rely="0.8", anchor="center")
            pbt16 = ttk.Button(fr5, text="2,8 €",  style="Basic.TButton", command=self.payment28).place(relx="0.3", rely="0.8", anchor="center")
            pbt17 = ttk.Button(fr5, text="3,0 €",  style="Basic.TButton", command=self.payment30).place(relx="0.4", rely="0.8", anchor="center")
            pbt18 = ttk.Button(fr5, text="3,2 €",  style="Basic.TButton", command=self.payment32).place(relx="0.5", rely="0.8", anchor="center")
            pbt19 = ttk.Button(fr5, text="3,4 €", style="Basic.TButton", command=self.payment34).place(relx="0.6", rely="0.8", anchor="center")
            pbt20 = ttk.Button(fr5, text="3,6 €",  style="Basic.TButton", command=self.payment36).place(relx="0.7", rely="0.8", anchor="center")
            pbt21 = ttk.Button(fr5, text="3,8 €",  style="Basic.TButton", command=self.payment38).place(relx="0.8", rely="0.8", anchor="center")
            #ttk.Label(fr5, text="", style="SuperEmpty.TLabel").pack()
            pbt22 = ttk.Button(fr6, text="4,0 €",  style="Basic.TButton", command=self.payment40).place(relx="0.2", rely="0.1", anchor="center")
            pbt23 = ttk.Button(fr6, text="4,2 €",  style="Basic.TButton", command=self.payment42).place(relx="0.3", rely="0.1", anchor="center")
            pbt24 = ttk.Button(fr6, text="4,4€",  style="Basic.TButton", command=self.payment44).place(relx="0.4", rely="0.1", anchor="center")
            pbt25 = ttk.Button(fr6, text="4,6 €",  style="Basic.TButton", command=self.payment46).place(relx="0.5", rely="0.1", anchor="center")
            pbt26 = ttk.Button(fr6, text="4,8 €", style="Basic.TButton", command=self.payment48).place(relx="0.6", rely="0.1", anchor="center")
            pbt27 = ttk.Button(fr6, text="5,0 €",  style="Basic.TButton", command=self.payment50).place(relx="0.7", rely="0.1", anchor="center")
            pbt28 = ttk.Button(fr6, text="5,2 €",  style="Basic.TButton", command=self.payment52).place(relx="0.8", rely="0.1", anchor="center")
            pbt29 = ttk.Button(fr6, text="5,4 €",  style="Basic.TButton", command=self.payment54).place(relx="0.4", rely="0.3", anchor="center")
            pbt30 = ttk.Button(fr6, text="5,6 €", style="Basic.TButton", command=self.payment56).place(relx="0.5", rely="0.3", anchor="center")
            pbt31 = ttk.Button(fr6, text="5,8 €",  style="Basic.TButton", command=self.payment58).place(relx="0.6", rely="0.3", anchor="center")
            pbt32 = ttk.Button(fr6, text="6,0 €",  style="Basic.TButton", command=self.payment60).place(relx="0.7", rely="0.3", anchor="center")
            ttk.Label(fr6, text="", style="SuperEmpty.TLabel").pack()
            ttk.Label(fr6, text="", style="SuperEmpty.TLabel").pack()
#------------------Define commands-------------------------------
    #button that does nothing!
    def donothing(self):
       filewin = Toplevel(gui)
       button = Button(filewin, text="Do nothing button")
       button.pack()
    #all payment setting
    def payment02(self):
        logging.debug('Payment 02 set')
        self.debt_to_add = float('{0:.2f}'.format(0.20))
        self.savepayment()

    def payment04(self):
        logging.debug('Payment 04 set')
        self.debt_to_add = float('{0:.2f}'.format(0.40))
        self.savepayment()

    def payment05(self):
        logging.debug('Payment 05 set')
        self.debt_to_add = float('{0:.2f}'.format(0.50))
        self.savepayment()

    def payment06(self):
        logging.debug('Payment 06 set')
        self.debt_to_add = float('{0:.2f}'.format(0.60))
        self.savepayment()

    def payment07(self):
        logging.debug('Payment 07 set')
        self.debt_to_add = float('{0:.2f}'.format(0.70))
        self.savepayment()

    def payment08(self):
        logging.debug('Payment 08 set')
        self.debt_to_add = float('{0:.2f}'.format(0.80))
        self.savepayment()

    def payment10(self):
        logging.debug('Payment 10 set')
        self.debt_to_add = float('{0:.2f}'.format(1.00))
        self.savepayment()

    def payment12(self):
        logging.debug('Payment 12 set')
        self.debt_to_add = float('{0:.2f}'.format(1.20))
        self.savepayment()

    def payment14(self):
        logging.debug('Payment 14 set')
        self.debt_to_add = float('{0:.2f}'.format(1.40))
        self.savepayment()

    def payment16(self):
        logging.debug('Payment 16 set')
        self.debt_to_add = float('{0:.2f}'.format(1.60))
        self.savepayment()

    def payment18(self):
        logging.debug('Payment 18 set')
        self.debt_to_add = float('{0:.2f}'.format(1.80))
        self.savepayment()

    def payment20(self):
        logging.debug('Payment 20 set')
        self.debt_to_add = float('{0:.2f}'.format(2.00))
        self.savepayment()

    def payment22(self):
        logging.debug('Payment 22 set')
        self.debt_to_add = float('{0:.2f}'.format(2.20))
        self.savepayment()

    def payment24(self):
        logging.debug('Payment 24 set')
        self.debt_to_add = float('{0:.2f}'.format(2.40))
        self.savepayment()

    def payment26(self):
        logging.debug('Payment 26 set')
        self.debt_to_add = float('{0:.2f}'.format(2.60))
        self.savepayment()

    def payment28(self):
        logging.debug('Payment 28 set')
        self.debt_to_add = float('{0:.2f}'.format(2.80))
        self.savepayment()

    def payment30(self):
        logging.debug('Payment 30 set')
        self.debt_to_add = float('{0:.2f}'.format(3.00))
        self.savepayment()

    def payment32(self):
        logging.debug('Payment 32 set')
        self.debt_to_add = float('{0:.2f}'.format(3.20))
        self.savepayment()

    def payment34(self):
        logging.debug('Payment 34 set')
        self.debt_to_add = float('{0:.2f}'.format(3.40))
        self.savepayment()

    def payment36(self):
        logging.debug('Payment 36 set')
        self.debt_to_add = float('{0:.2f}'.format(3.60))
        self.savepayment()

    def payment38(self):
        logging.debug('Payment 38 set')
        self.debt_to_add = float('{0:.2f}'.format(3.80))
        self.savepayment()

    def payment40(self):
        logging.debug('Payment 40 set')
        self.debt_to_add = float('{0:.2f}'.format(4.00))
        self.savepayment()

    def payment42(self):
        logging.debug('Payment 42 set')
        self.debt_to_add = float('{0:.2f}'.format(4.20))
        self.savepayment()

    def payment44(self):
        logging.debug('Payment 44 set')
        self.debt_to_add = float('{0:.2f}'.format(4.40))
        self.savepayment()

    def payment46(self):
        logging.debug('Payment 46 set')
        self.debt_to_add = float('{0:.2f}'.format(4.60))
        self.savepayment()

    def payment48(self):
        logging.debug('Payment 48 set')
        self.debt_to_add = float('{0:.2f}'.format(4.80))
        self.savepayment()

    def payment50(self):
        logging.debug('Payment 50 set')
        self.debt_to_add = float('{0:.2f}'.format(5.00))
        self.savepayment()

    def payment52(self):
        logging.debug('Payment 52 set')
        self.debt_to_add = float('{0:.2f}'.format(5.20))
        self.savepayment()

    def payment54(self):
        logging.debug('Payment 54 set')
        self.debt_to_add = float('{0:.2f}'.format(5.40))
        self.savepayment()

    def payment56(self):
        logging.debug('Payment 56 set')
        self.debt_to_add = float('{0:.2f}'.format(5.60))
        self.savepayment()

    def payment58(self):
        logging.debug('Payment 58 set')
        self.debt_to_add = float('{0:.2f}'.format(5.80))
        self.savepayment()

    def payment60(self):
        logging.debug('Payment 60 set')
        self.debt_to_add = float('{0:.2f}'.format(6.00))
        self.savepayment()

#---------START OF DEV TOOLS SECTION--------------------------------------------------------

    def devtool_gui(self):
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        #import settings
        self.devwin = Toplevel(gui)
        self.devwin.geometry("+{}+{}".format(positionRight, positionDown))
        #Set GUI parameters - Title
        self.devwin.title("Debug")
        #Set GUI parameters - Window size x+y
        #self.devwin.geometry("330x300")
        #frame
        dev_fr1=ttk.Frame(self.devwin,relief="raised", borderwidth=1)
        dev_fr1.pack(expand=0, fill="both")

        #style
        devstyle = ttk.Style()
        devstyle.configure("Devinfo.TLabel", font=("calibri", 20, "bold"))
        devstyle.configure("Devtxt.TLabel", font=("calibri", 15))

        dev_lbl1 = ttk.Label(dev_fr1, text="Debug tools", style="Devinfo.TLabel").pack()
        self.dev_infotxt = StringVar()
        self.dev_infotxt.set('Info')
        dev_lbl2 = ttk.Label(dev_fr1, textvariable=self.dev_infotxt,style="Devtxt.TLabel").pack()
        #Check buttons
        #empty label
        ttk.Label(dev_fr1, text="").pack()
        btn1 = ttk.Button(dev_fr1, text='Delete logs', command=self.deletelogs).pack()
        #empty label
        ttk.Label(dev_fr1, text="").pack()
        btn2 = ttk.Button(dev_fr1, text='Add random payments', command=self.devaddpayments).pack()
        #empty label
        ttk.Label(dev_fr1, text="").pack()
        btn3 = ttk.Button(dev_fr1, text='Check db size', command=self.check_db_size).pack()
        #empty label
        ttk.Label(dev_fr1, text="").pack()
        #Save setting button
        btn4 = ttk.Button(dev_fr1, text='Endurance test', command=self.endurance_test).pack()
        #Quit button
        btn5 = ttk.Button(dev_fr1, text='Back', command=self.devwin.destroy).pack()
        #empty label
        ttk.Label(dev_fr1, text="").pack()

    def deletelogs(self):

        dir_path= str(os.path.dirname(__file__))
        folder = dir_path+'\log'
        #print (folder)
        try:
            for filename in os.listdir(folder):
                file_path = os.path.join(folder, filename)

                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)

            self.dev_infotxt.set('Logs deleted')
            logging.info('*devtool* Logs deleted')
        except Exception as e:
            logging.error('*devtool* Failed to delete logs %s. Reason: %s' % (file_path, e))
            self.dev_infotxt.set('-ERROR-')

    def devaddpayments(self):

        names_import = self.test_list
        self.plist = list(names_import)
        self.srtplist = sorted(self.plist)
        for item in self.srtplist:

            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            logging.debug('*devaddpayments* Connecting to DB people.db')
            logging.info('*devaddpayments* Adding payments to DB')

            c.execute('SELECT name FROM people WHERE name=?', (item,))
            name_call = c.fetchall()
            name_str = str(name_call)
            name_str = name_str.replace('[','')
            name_str = name_str.replace(']','')
            name_str = name_str.replace('(','')
            name_str = name_str.replace(')','')
            name_str = name_str.replace(',','')
            name_str = name_str.replace('\'','')
            paytime = strftime('%H:%M:%S')
            paydate = strftime("%d-%m-%Y")
            value = randint(0,100)
            #payname = ''
            #payname = items
            c.execute('UPDATE people SET date=?, time=?, debt=? WHERE name=?',
            (paydate , paytime , value, name_str,))
            print("Adding debt to:")
            print(name_str)
            print(value)

            #commit the changes to db
            conn.commit()
            #close the connection
            conn.close()
            #close the connection
            datetime = strftime("%Y-%m-%d %H:%M:%S")
            conn = sqlite3.connect('history\\'+self.debt_history)
            c = conn.cursor()
            logging.debug('*devaddpayments* Connecting to DB debts')
            self.infotext.set("Connected to DB")
            sqlite_insert_with_param = """INSERT INTO debts
                                (datetime, name, debt)
                                VALUES (?, ?, ?);"""
            data_tuple = (datetime, name_str , value)

            #insert row
            c.execute(sqlite_insert_with_param, data_tuple)

            self.infotext.set("Writing to DB done, closing")
            logging.debug('*devaddpayments* Closing DB debts')
            #commit the changes to db
            conn.commit()
            #close the connection
            conn.close()
            #END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!
            logging.info('*devaddpayments* Payment backup saved to DB')

        logging.info('*devaddpayments* Added values to DB people')
        self.dev_infotxt.set('Maksu tallennettu')

    def check_db_size(self):

        dir_path = str(os.path.dirname(__file__))
        hist_folder = dir_path + '\history\\'
        log_folder = dir_path + '\log\\'
        db_folder = dir_path + '\DB\\'
        print(hist_folder)

        # initialize the size
        db_size = 0
        # use the walk() method to navigate through directory tree
        for dirpath, dirnames, filenames in os.walk(db_folder):
            for i in filenames:
                # use join to concatenate all the components of path
                f = os.path.join(dirpath, i)

                # use getsize to generate size in bytes and add it to the total size
                db_size += os.path.getsize(f)
        db_size_mb = db_size/1048576
        print('DB size (MB):')
        print(db_size_mb)

        # initialize the size
        hist_size = 0
        # use the walk() method to navigate through directory tree
        for dirpath, dirnames, filenames in os.walk(hist_folder):
            for i in filenames:
                # use join to concatenate all the components of path
                f = os.path.join(dirpath, i)

                # use getsize to generate size in bytes and add it to the total size
                hist_size += os.path.getsize(f)
        hist_size_mb = hist_size / 1048576
        print('History DB size (MB):')
        print(hist_size_mb)

        # initialize the size
        log_size = 0
        # use the walk() method to navigate through directory tree
        for dirpath, dirnames, filenames in os.walk(log_folder):
            for i in filenames:
                # use join to concatenate all the components of path
                f = os.path.join(dirpath, i)

                # use getsize to generate size in bytes and add it to the total size
                log_size += os.path.getsize(f)
        log_size_mb = log_size / 1048576
        print('Log DB size (MB):')
        print(log_size_mb)

    def endurance_test(self):
        i = 0
        count = input("Number of payments: ")
        count_int = int(count)
        sleep_count = input("Sleep time between inserts: ")
        sleep_int = int(sleep_count)
        while i < count_int:
            count = str(i)
            print('Count: '+ count)
            set_name = random.choice(self.test_list)
            print('Name:' + set_name)
            self.getname.set(set_name)
            func_list = [self.payment02, self.payment04, self.payment06, self.payment08, self.payment10, self.payment12, self.payment14, self.payment16, self.payment18, self.payment20, self.payment22, self.payment24, self.payment26, self.payment28, self.payment30, self.payment32, self.payment34, self.payment36, self.payment38, self.payment40, self.payment42, self.payment44, self.payment46, self.payment48, self.payment50, self.payment52, self.payment54, self.payment56, self.payment58, self.payment60]
            random.choice(func_list)()
            print('Sleep:')
            print(sleep_int)
            sleep(sleep_int)
            i += 1

#---------END OF DEV TOOLS SECTION--------------------------------------------------------
    def nm_mngmt_pswrd(self):
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        self.nm_mngmt_pswrd = Toplevel(gui)
        self.nm_mngmt_pswrd.geometry("+{}+{}".format(positionRight, positionDown))
        self.nm_mngmt_pswrd.attributes("-topmost", True)
        nm_mngmt_fr = ttk.Frame(self.nm_mngmt_pswrd, relief="flat", borderwidth=1)
        nm_mngmt_fr.pack(expand=0, fill="both")
        ttk.Label(nm_mngmt_fr, text="Password:").pack()
        self.dev_ent = ttk.Entry(nm_mngmt_fr, justify = 'center')
        self.dev_ent.pack()
        btn1 = ttk.Button(nm_mngmt_fr, text='Ok', command=self.nm_mngmt_checkpswrd).pack()
        btn2 = ttk.Button(nm_mngmt_fr, text='Cancel', command=self.nm_mngmt_pswrd.destroy).pack()

    def nm_mngmt_checkpswrd(self):
        input = str(self.dev_ent.get())
        if input == 'Kahvikassa!':
            logging.info('*nm_mngmt* Correct password, opening GUI')
            self.nm_mngmt_pswrd.destroy()
            self.nm_mngmt()
        elif input == 'Debug!':
            logging.info('*nm_mngmt* Correct password, opening Debug GUI')
            self.nm_mngmt_pswrd.destroy()
            self.devtool_gui()
        else:
            logging.info('*nm_mngmt* Incorrect password, closing')
            self.nm_mngmt_pswrd.destroy()

    def nm_mngmt(self):
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        global name_list
        logging.info('*nm_mngmt* GUI Starting')
        #name_management version - V03
        self.nm_mngmt_gui = Toplevel(gui)
        self.nm_mngmt_gui.geometry("+{}+{}".format(positionRight, positionDown))
        self.nm_mngmt_gui.attributes("-topmost", True)
        fr1=ttk.Frame(self.nm_mngmt_gui,relief="raised", borderwidth=1)
        fr1.pack(expand=0, fill="both")

        fr2=ttk.Frame(self.nm_mngmt_gui,relief="raised", borderwidth=0)
        fr2.pack(expand=0, fill="both")

        fr3=ttk.Frame(self.nm_mngmt_gui,relief="raised", borderwidth=0)
        fr3.pack(expand=0, fill="both")
        #---------------------Define styling options ----------------------------------
        #Label
        style = ttk.Style()
        style.configure("NMMain.TLabel", font=("calibri", 13, "bold"))
        style.configure("NMLbox.TLabel", font=("calibri", 11, "bold"))
        style.configure("Empty.TLabel", font=("calibri", 8, "bold"))
        style.configure("NMBasic.TButton", font=("calibri", 12, "bold"))
        #---------------------Define variables ----------------------------------
        self.nm_infotext = StringVar("")
        self.nm_infotext.set("")
        self.DB_changed = False

        #-----------Connect to SQLite table and create table if not existing------------
        try:
           conn = sqlite3.connect('DB\people.db')
           c = conn.cursor()
           logging.debug('*nm_mngmt* Connecting to DB people')
           self.nm_infotext.set("Connecting to DB")

           c.execute("SELECT name FROM people")
           rows = c.fetchall()
           name_list = []
           for row in rows:
               name_list += row

           self.nm_infotext.set("DB read")
           #commit the changes to db
           conn.commit()
           #close the connection
           conn.close()
           logging.debug('*nm_mngmt* Closed DB people')

        except Exception as e:
           logging.error('*nm_mngmt* -ERROR- Reason: %s' % (e))
           self.nm_infotext.set("-ERROR-")
           username = os.getlogin()  # Fetch username
           date = strftime("%d-%m-%Y")
           print('Creating error log')
           error_to_write = '*nm_mngmt* -ERROR- Reason: %s' % (e)
           f = open(f"C:\\Users\\{username}\\Desktop\\"+date+"_error_info.txt", "w")
           f.write(ver)
           f.write('\n')
           f.write(error_to_write)
           f.write('\n')
           f.close()
           error_info = str(date+'_error_info.txt')
           messagebox.showerror("Error", "Something went wrong - error description ("+error_info+") has been created to desktop")

       #---------------------Main pm_gui ----------------------------------
        self.lbl = ttk.Label(fr1, text="Henkilöstö", style="NMMain.TLabel").pack()
        self.lbl1 = ttk.Label(fr1, text='Info',style="NMMain.TLabel").pack()
        self.lbl2 = ttk.Label(fr1, textvariable=self.nm_infotext,style="NMLbox.TLabel").pack()

        self.plist = list(name_list)
        self.srtplist = sorted(self.plist)
        #self.f.close()

        self.scrollbar = Scrollbar(fr2)
        self.scrollbar.pack(side=RIGHT, fill=Y)

        self.listbox = Listbox(fr2, yscrollcommand=self.scrollbar.set)
        self.listbox.pack()
        self.scrollbar.config(command=self.listbox.yview)

        for item in self.srtplist:
           self.listbox.insert(END, item)

        self.lb = self.listbox
        self.elbl1 = ttk.Label(fr3, text='',style="Empty.TLabel").pack()
        self.delbtn = ttk.Button(fr3, text="Poista nimi", style="NMBasic.TButton",
                  command=self.delp) #lb=self.lb: lb.delete(ANCHOR))
        self.delbtn.pack()
        self.elbl2 = ttk.Label(fr3, text='',style="Empty.TLabel").pack()
        #self.lbl3 = ttk.Label(fr3, text="Add name", style="NMMain.TLabel").pack()
        self.ent = ttk.Entry(fr3)
        self.ent.pack()
        self.addbtn = ttk.Button(fr3, text="Lisää nimi", style="NMBasic.TButton",
                  command=self.savep)
        self.addbtn.pack()
        self.elbl3 = ttk.Label(fr3, text='',style="Empty.TLabel").pack()
        self.imbtn = ttk.Button(fr3, text="Tuo nimiä listaan", style="NMBasic.TButton",
                  command=self.nm_mngmt_import)
        self.imbtn.pack()
        self.expbtn = ttk.Button(fr3, text="Tallenna nimilista", style="NMBasic.TButton",
                  command=self.nm_mngmt_export)
        self.expbtn.pack()
        self.elbl4 = ttk.Label(fr3, text='',style="Empty.TLabel").pack()
        self.qbtn = ttk.Button(fr3, text="Takaisin", style="NMBasic.TButton",
                  command=self.close_nm_mngmt)
        self.qbtn.pack()
        self.elbl4 = ttk.Label(fr3, text='',style="Empty.TLabel").pack()

    def nm_mngmt_export(self):
        try:
            month_year = strftime("%d-%m-%Y")
            f = io.open(month_year+"_nimilista.list", "w", encoding='utf8')
            list = []
            for row in self.srtplist:
                to_list = row.replace('\n', '')
                print(to_list)
                list.append(to_list)

            towrite = str(list)
            print(towrite)
            f.write(towrite)
            f.close()
            self.nm_infotext.set("Nimilista tallennettu")
            logging.debug('*nm_mngmt* Names exported')

        except Exception as e:
            logging.error('*nm_mngmt_export* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*nm_mngmt* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error","Something went wrong - description (" + error_info + ") created to desktop")
            self.nm_infotext.set("-ERROR-")

    def nm_mngmt_import(self):
        self.top = Toplevel(gui)
        #top.geometry("150x150")
        self.im_lbl1_vrbl = StringVar()
        self.im_lbl1_vrbl.set('')
        import_fr1=ttk.Frame(self.top,relief="raised", borderwidth=1)
        import_fr1.pack(expand=1, fill="both")

        import_fr2=ttk.Frame(self.top,relief="raised", borderwidth=1)
        import_fr2.pack(expand=1, fill="both")
        import_fr2_1=ttk.Frame(import_fr2,relief="raised", borderwidth=1)
        import_fr2_1.grid(column='0', row='0')
        import_fr2_2=ttk.Frame(import_fr2,relief="raised", borderwidth=1)
        import_fr2_2.grid(column='1', row='0')

        import_fr3=ttk.Frame(self.top,relief="raised", borderwidth=1)
        import_fr3.pack(expand=1, fill="both")

        import_fr4=ttk.Frame(self.top,relief="raised", borderwidth=1)
        import_fr4.pack(expand=1, fill="both")
        ttk.Label(import_fr1,text="Tuo nimilista", style = "Mlabel.TLabel").pack()
        im_lbl = ttk.Label(import_fr2_1,textvariable=self.im_lbl1_vrbl,width=80 ,style = "Browse.TLabel").grid(column='0', row='0')
        ttk.Button(import_fr2_2, text='Etsi', command=self.fileDialog, style = "Btn.TButton").grid(column='1', row='0')

        ttk.Button(import_fr3, text='OK', command=self.importnames, style = "Btn.TButton").pack(anchor='e')
        ttk.Button(import_fr3, text='Takaisin', command=self.top.destroy, style = "Btn.TButton").pack(anchor='e')

    def fileDialog(self):
        try:
            self.filename = filedialog.askopenfilename(initialdir =  "/", title = "Select A File", filetype =
            (("Name list files","*.list"),("all files","*.*")) )
            self.im_lbl1_vrbl.set(self.filename)
            f = io.open(self.filename, "r", encoding='utf8')
            self.file_str = f.read()
            f.close()
        except Exception as e:
            logging.error('*fileDialog* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*fileDialog* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - error description (" + error_info + ") has been created to desktop")

    def importnames(self):
        try:

            names_import = ast.literal_eval(self.file_str)
            #print(self.file_str)
            #self.plist = list(self.file_str)
            #self.srtplist = sorted(self.plist)

            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            logging.debug('*nm_mngmt* Connecting to DB people.db')
            logging.info('*nm_mngmt* Importing namelist to DB')

            for x in names_import:
                print(x)
                c.execute('SELECT name FROM people WHERE name=?', (x,))
                name_call = c.fetchall()
                name_str = str(name_call)

                if name_str != x:
                    logging.debug('*nm_mngmt* Name not found adding to DB')
                    self.listbox.insert(END, x)
                    debtint = 0.0
                    time = strftime('%H:%M:%S')
                    date = strftime("%d-%m-%Y")
                    name = x
                    #payname = x.rstrip("\n")

                    sqlite_insert_with_param = """INSERT INTO people
                                          (date, time, name, debt)
                                          VALUES (?, ?, ?, ?);"""
                    data_tuple = (date , time , name , debtint)

                    #insert line
                    c.execute(sqlite_insert_with_param, data_tuple)
                    #commit the changes to db
                    conn.commit()
                    self.DB_changed = True

                    logging.info('*nm_mngmt* Name added')
                    logging.info(name)

                else:
                    logging.debug('*nm_mngmt* Name already in DB, pass')
                    pass
            #close the connection
            conn.close()
            logging.debug('*nm_mngmt* Close DB people')
            self.lb = self.listbox
            self.top.destroy()
            logging.info('*nm_mngmt* Namelist imported to DB')
            self.nm_infotext.set("Nimilista tuotu")

        except Exception as e:
            logging.error('*importnames* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*importnames* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.nm_infotext.set("-ERROR-")

    def close_nm_mngmt(self):

        if self.DB_changed == True:
            self.nm_mngmt_gui.destroy()
            logging.info('*nm_mngmt* Database updated, restarting program...')
            self.DB_changed = False
            messagebox.showwarning(title='Varoitus!', message='Nimilista päivitetty, ohjelma käynnistetään uudelleen.')
            #os.execl(sys.executable, sys.executable, *sys.argv)
            os.execl(sys.executable, os.path.abspath(__file__), *sys.argv)
            print("Database updated, restarting program...")
        else:
            logging.debug('*nm_mngmt* No changes -> self.nm_mngmt_gui.destroy')
            self.nm_mngmt_gui.destroy()

    def delp(self):
        try:
            self.selname = self.listbox.get(ACTIVE)
            self.listbox.delete(ACTIVE)
            self.plist.remove(self.selname)
            # start of SQLite code for deleting person
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            self.nm_infotext.set("Connecting to DB")
            logging.debug('*nm_mngmt* Connecting to DB people')
            # delete row with specific name from table
            c.execute('DELETE FROM people WHERE name=?', (self.selname,))
            logging.info('*nm_mngmt* Name deleted:')
            logging.info(self.selname)

            #commit the changes to db
            conn.commit()
            #close the connection
            conn.close()
            logging.debug('*nm_mngmt* Closed DB people.db')
            self.DB_changed = True
            self.nm_infotext.set("Nimi poistettu")

        except Exception as e:
            logging.error('*delp* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*delp* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - error (" + error_info + ") created to desktop")
            self.nm_infotext.set("-ERROR-")

    def savep(self):
        try:
            input = self.ent.get()
            self.listbox.insert(END, input)
            debtint = 0.00
            time = strftime('%H:%M:%S')
            date = strftime("%d-%m-%Y")
            name = str(input)
            name = name.rstrip("\n")
            name = name.strip()
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            logging.debug('*nm_mngmt* Connecting to DB people')
            sqlite_insert_with_param = """INSERT INTO people
                                  (date, time, name, debt)
                                  VALUES (?, ?, ?, ?);"""
            data_tuple = (date , time , name , debtint)

            #insert row
            c.execute(sqlite_insert_with_param, data_tuple)
            #commit the changes to db
            conn.commit()
            #close the connection
            conn.close()
            logging.debug('*nm_mngmt* Close DB people')
            self.nm_infotext.set("Nimi lisätty!")
            self.ent.delete(0, 'end')
            self.DB_changed = True
            logging.info('*nm_mngmt* Name added')
            logging.info(name)

        except Exception as e:
            logging.error('*savep* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*savep* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.nm_infotext.set("-ERROR-")

    def savepayment(self):
        try:
            #Writes payment into independent DB -> debts
            paytime = strftime('%H:%M:%S')
            paydate = strftime("%d-%m-%Y")

            payname = self.getname.get()
            logging.debug('*savepayment* Check if name is in DB')
            #Checks if given name is in people DB
            logging.debug('*savepayment* Connecting to DB people')
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()
            #print("Connected to SQLite DB people")
            self.infotext.set("Connected to DB")

            c.execute('SELECT name FROM people WHERE name=?', (payname,))
            name_db = c.fetchall()
            name_check = str(name_db)

            name_check = name_check.replace('\\','')
            name_check = name_check.replace('[','')
            name_check = name_check.replace(']','')
            name_check = name_check.replace('(','')
            name_check = name_check.replace(')','')
            name_check = name_check.replace(',','')
            name_check = name_check.replace('\'','')

            self.infotext.set("DB read, closing DB")
            conn.commit()
            #close the connection
            conn.close()
            logging.debug('*savepayment* Closing DB people')

            if payname == '': #Checks if namebox is empty
                logging.debug('*savepayment* Name check - Empty cell')
                self.infotext.set("Syötä nimi!")

            elif payname != name_check: #Checks if given name is in people DB
                logging.debug('*savepayment* Name check - Name not found')
                messagebox.showerror("Virhe!", "Nimeä "+payname+" ei löydy!\n" + 'Tarkista että nimi on oikein...')
                self.getname.set('')
                self.infotext.set('')
            else:
                date_now = strftime("%m-%Y")
                db_path = date_now + "_debts.db"
                if date_now != self.debt_history:
                    logging.info('Month change detected, creating a new history DB')
                    self.debt_history = db_path
                    logging.info(self.debt_history)
                    conn = sqlite3.connect('history\\' + self.debt_history)
                    c = conn.cursor()
                    # create table 'people'
                    c.execute('''CREATE TABLE IF NOT EXISTS debts
                                             (datetime, name, debt, payment)''')

                    conn.commit()
                    # close the connection
                    conn.close()
                    logging.info('New history DB: ' + self.debt_history)
                    self.infotext.set("New history DB created")
                    # Writes payment into separate backup DB --- this is solely for backup purposes only!!!!
                    datetime = strftime("%Y-%m-%d %H:%M:%S")
                    conn = sqlite3.connect('history\\' + self.debt_history)
                    c = conn.cursor()
                    logging.debug('*savepayment* Connecting to DB debt history')
                    self.infotext.set("Connected to DB")
                    sqlite_insert_with_param = """INSERT INTO debts
                                                            (datetime, name, debt)
                                                            VALUES (?, ?, ?);"""
                    data_tuple = (datetime, payname, self.debt_to_add)

                    # insert row
                    c.execute(sqlite_insert_with_param, data_tuple)

                    self.infotext.set("Writing to DB done, closing")
                    logging.debug('*savepayment* Closing DB' + self.debt_history)
                    # commit the changes to db
                    conn.commit()
                    # close the connection
                    conn.close()
                    # END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!
                    logging.info('*savepayment* Payment backup saved to DB')
                    # Writes payment into people DB
                    conn = sqlite3.connect('DB\people.db')
                    c = conn.cursor()
                    logging.debug('*savepayment* Connecting to DB people')
                    self.infotext.set("Connected to DB")

                    c.execute("SELECT debt FROM people WHERE name=?", (payname,))
                    debt = c.fetchall()
                    # print(debt)
                    debt_tuple = []
                    debt_float = 0.00
                    for row in debt:
                        debt_tuple += row
                        for row in debt_tuple:
                            debt_float += float(row)
                            debt_float = round(debt_float,1)
                            #print('Debt_float:')
                            #print(debt_float)

                    # print("Read debt from DB people, closing")
                    self.infotext.set("Read debt from DB, closing")
                    logging.debug('*savepayment* Existing debts read from DB')
                    old_debt = str(debt_float)
                    old_debt = old_debt[0:5]
                    # print('old_debt:')
                    # print(old_debt)
                    debts_conc = float(old_debt) + self.debt_to_add
                    # print('debts_conc:')
                    # print(debts_conc)
                    debt_to_save_clnd = str(debts_conc)
                    debt_to_save_clnd = debt_to_save_clnd[0:5]
                    debt_to_save = float(debt_to_save_clnd)
                    # print('debts_to_save:')
                    # print(debt_to_save)

                    c.execute('UPDATE people SET date=?, time=?, debt=? WHERE name=?',
                              (paydate, paytime, debt_to_save, payname))
                    self.infotext.set("Maksu kirjattu")
                    self.infotext2.set(paytime)
                    logging.debug('*savepayment* Name and debt to save')
                    logging.debug(payname)
                    logging.debug(debt_to_save)
                    # commit the changes to db
                    conn.commit()
                    # close the connection
                    conn.close()
                    logging.info('*savepayment* Update succesfull, closing DB people')
                    # clear namebox
                    self.getname.set('')
                    self.name_for_msgbox = payname
                    self.time_for_msgbox = paytime
                    # messagebox.showinfo(title='Maksu kirjattu', message=payname+' ,maksu '+ debt_to_msgbox +"€ kirjattu")
                    self.paywin()

                else:

                    #Writes payment into separate backup DB --- this is solely for backup purposes only!!!!
                    datetime = strftime("%Y-%m-%d %H:%M:%S")
                    conn = sqlite3.connect('history\\'+self.debt_history)
                    c = conn.cursor()
                    logging.debug('*savepayment* Connecting to DB debt history')
                    self.infotext.set("Connected to DB")
                    sqlite_insert_with_param = """INSERT INTO debts
                                        (datetime, name, debt)
                                        VALUES (?, ?, ?);"""
                    data_tuple = (datetime, payname , self.debt_to_add)

                    #insert row
                    c.execute(sqlite_insert_with_param, data_tuple)

                    self.infotext.set("Writing to DB done, closing")
                    logging.debug('*savepayment* Closing DB'+self.debt_history)
                    #commit the changes to db
                    conn.commit()
                    #close the connection
                    conn.close()
                    #END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!
                    logging.info('*savepayment* Payment backup saved to DB')
                    #Writes payment into people DB
                    conn = sqlite3.connect('DB\people.db')
                    c = conn.cursor()
                    logging.debug('*savepayment* Connecting to DB people')
                    self.infotext.set("Connected to DB")

                    c.execute("SELECT debt FROM people WHERE name=?", (payname,))
                    debt = c.fetchall()
                    #print(debt)
                    debt_tuple = []
                    debt_float = 0.00
                    for row in debt:
                        debt_tuple += row
                        for row in debt_tuple:
                            debt_float += float(row)
                            #print('Debt_float:')
                            #print(debt_float)

                    #print("Read debt from DB people, closing")
                    self.infotext.set("Read debt from DB, closing")
                    logging.debug('*savepayment* Existing debts read from DB')
                    old_debt = str(debt_float)
                    old_debt = old_debt[0:5]
                    #print('old_debt:')
                    #print(old_debt)
                    debts_conc = float(old_debt) + self.debt_to_add
                    #print('debts_conc:')
                    #print(debts_conc)
                    debt_to_save_clnd = str(debts_conc)
                    debt_to_save_clnd = debt_to_save_clnd[0:5]
                    debt_to_save = float(debt_to_save_clnd)
                    #print('debts_to_save:')
                    #print(debt_to_save)

                    c.execute('UPDATE people SET date=?, time=?, debt=? WHERE name=?', (paydate , paytime , debt_to_save, payname))
                    self.infotext.set("Maksu kirjattu")
                    self.infotext2.set(paytime)
                    logging.debug('*savepayment* Name and debt to save')
                    logging.debug(payname)
                    logging.debug(debt_to_save)
                    #commit the changes to db
                    conn.commit()
                    #close the connection
                    conn.close()
                    logging.info('*savepayment* Update succesfull, closing DB people')
                    #clear namebox
                    self.getname.set('')
                    self.name_for_msgbox = payname
                    self.time_for_msgbox = paytime
                    #messagebox.showinfo(title='Maksu kirjattu', message=payname+' ,maksu '+ debt_to_msgbox +"€ kirjattu")
                    self.paywin()

        except Exception as e:
            logging.error('*savepayment* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*savepayment* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.infotext.set("-ERROR-")

    def paywin(self):
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        #print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        self.debt_to_msgbox = str(self.debt_to_add)
        self.paywin_gui = Toplevel(gui)
        # Positions the window in the center of the page.
        self.paywin_gui.geometry("+{}+{}".format(positionRight, positionDown))
        self.paywin_gui.attributes("-topmost", True)
        paywin_fr = ttk.Frame(self.paywin_gui, relief="flat", borderwidth=1)
        paywin_fr.pack(expand=0, fill="both")
        ttk.Label(paywin_fr, text=self.time_for_msgbox).pack()
        ttk.Label(paywin_fr, text=self.name_for_msgbox).pack()
        ttk.Label(paywin_fr, text="Maksu " + self.debt_to_msgbox + " € kirjattu.").pack()
        btn1 = ttk.Button(paywin_fr, text='Ok', command=self.paywin_gui.destroy).pack()
        self.paywin_gui.after(10000, lambda: self.paywin_gui.destroy())  # Destroy the widget after 30 seconds

    def debtquery(self):
        try:

            name = self.getname.get()
            qrname = name.rstrip("\n")
            if qrname == '':
                logging.debug('*debtquery* No name given')
                self.infotext.set("Syötä nimi!")

            else:
                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()
                self.infotext.set("Connected to DB")
                logging.debug('*debtquery* Connected to DB people')
                query_data = (qrname)
                c.execute('SELECT debt FROM people WHERE name=?', (qrname,))
                curr_debt = c.fetchall()

                curr_debt_tuple = []
                curr_debt_str = ''
                for row in curr_debt:
                    curr_debt_tuple += row
                    for row in curr_debt_tuple:
                        curr_debt_str += str(row)
                        #print(debt_float)
                cleaned_debt_str = curr_debt_str[0:5]
                self.infotext.set("Debt read from DB")
                logging.debug('*debtquery* Read debt from DB, closing DB people')
                #commit the changes to db
                conn.commit()
                #close the connection
                conn.close()

                #Create a messagebox for debt amount
                messagebox.showinfo(qrname, "Velan määrä: \n" + cleaned_debt_str +'€')
                #print("Closing SQLite DB people")

                self.infotext.set('')
                self.getname.set('')

        except Exception as e:
            logging.error('*debtquery* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*debtquery* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.infotext.set("-ERROR-")



    def about(self):
        #filewin = ThemedTk(theme="arc")
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        aboutwin = Toplevel(gui)
        aboutwin.geometry("+{}+{}".format(positionRight, positionDown))
        aboutfr=ttk.Frame(aboutwin,relief="raised", borderwidth=1)
        aboutfr.pack(expand=0, fill="both")

        aboutwin.title("About")
        #Set GUI parameters - Is window resizable
        aboutwin.resizable(False, False)
        ttk.Label(aboutfr, text="Kahvikassa ohjelma", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="Versio: "+ver, style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="Machine ID:", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text=self.curr_mID, style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="Copyright 2020", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="Juha Rouvinen", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="jmrouvinen@gmail.com", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="All rights reserved.", style="Other.TLabel").pack()
        ttk.Label(aboutfr, text='Do not alter the code, reverse engineer', style="Other.TLabel").pack()
        ttk.Label(aboutfr, text='or distribute the software', style="Other.TLabel").pack()
        ttk.Label(aboutfr, text='without the copyright owners consent.', style="Other.TLabel").pack()
        ttk.Label(aboutfr, text="", style="Other.TLabel").pack()
        button = Button(aboutfr, text="Back", command=aboutwin.destroy)
        button.pack()

    def helpindx(self):
        filewin = Toplevel(gui)
        button = Button(filewin, text="Do nothing button")
        button.pack()
#-------------- START OF DEBT MANAGEMENT ----------------------------------------
    def debtmang_pswrd(self):
        # filewin = ThemedTk(theme="arc")
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        self.debtmang_pswrd = Toplevel(gui)
        self.debtmang_pswrd.geometry("+{}+{}".format(positionRight, positionDown))
        self.debtmang_pswrd.attributes("-topmost", True)
        debtmang_fr = ttk.Frame(self.debtmang_pswrd, relief="flat", borderwidth=1)
        debtmang_fr.pack(expand=0, fill="both")
        ttk.Label(debtmang_fr, text="Password:").pack()
        self.dev_ent = ttk.Entry(debtmang_fr, justify = 'center')
        self.dev_ent.pack()
        btn1 = ttk.Button(debtmang_fr, text='Ok', command=self.debtmang_checkpswrd).pack()
        btn2 = ttk.Button(debtmang_fr, text='Cancel', command=self.debtmang_pswrd.destroy).pack()

    def debtmang_checkpswrd(self):
        input = str(self.dev_ent.get())
        if input == 'Kahvikassa!':
            logging.info('*debtmang* Correct password, opening GUI')
            self.debtmang_pswrd.destroy()
            self.debtmang()

        else:
            logging.info('*debtmang* Incorrect password, closing')
            self.filewin.destroy()

    def debtmang(self):

        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        debtmang_gui = Toplevel(gui)
        debtmang_gui.geometry("+{}+{}".format(positionRight, positionDown))
        debtmang_gui.attributes("-topmost", True)
        logging.info('*debtmang* GUI Starting')
        #debtmang_gui.geometry("200x1500")
        #frames
        debtmang_fr1=ttk.Frame(debtmang_gui,relief="raised", borderwidth=1)
        debtmang_fr1.pack(expand=0, fill="both")

        debtmang_fr2=ttk.Frame(debtmang_gui,relief="raised", borderwidth=0)
        debtmang_fr2.pack(expand=0, fill="both")

        debtmang_fr3=ttk.Frame(debtmang_gui,relief="raised", borderwidth=0)
        debtmang_fr3.pack(expand=0, fill="both")
        #---------------------Define styling options ----------------------------------
        #Labels
        style = ttk.Style()
        style.configure("DMMain.TLabel", font=("calibri", 13, "bold"))
        style.configure("DMLbox.TLabel", font=("calibri", 11, "bold"))
        style.configure("Empty.TLabel", font=("calibri", 3, "bold"))
        style.configure("DMBasic.TButton", font=("calibri", 12, "bold"))
        #---------------------Define GUI settings END------------------------------------
        self.DM_infotext = StringVar("")
        self.DM_infotext.set("")
        self.DM_infotext2 = StringVar("")
        self.DM_infotext2.set("")
        self.EP_infotext = StringVar("")
        self.EP_infotext.set("")
        self.EP_infotext2 = StringVar("")
        self.EP_infotext2.set("")
        #Create connection into DB and get names and debts
        conn = sqlite3.connect('DB\people.db')
        c = conn.cursor()

        logging.info('*debtmang* Connected to DB people')
        self.DM_infotext.set("Connected to DB")

        c.execute('SELECT name, debt FROM people ORDER BY name ASC;')
        names_debts = c.fetchall()

        conn.commit()
        #close the connection
        conn.close()

        #empty label
        ttk.Label(debtmang_fr1, text="").pack()

        self.debtmang_tree_header = ['Nimi', 'Velka']

        self.debtmang_tree = ttk.Treeview(debtmang_fr2, columns=self.debtmang_tree_header, show="headings")
        vsb = ttk.Scrollbar(debtmang_fr2, orient="vertical",
            command=self.debtmang_tree.yview)
        hsb = ttk.Scrollbar(debtmang_fr2, orient="horizontal",
            command=self.debtmang_tree.xview)
        self.debtmang_tree.configure(yscrollcommand=vsb.set,
            xscrollcommand=hsb.set)
        self.debtmang_tree.bind("<Double-Button>", self.editpayment)
        self.debtmang_tree.grid(column=0, row=0, sticky='nsew',)
        vsb.grid(column=1, row=0, sticky='ns')
        hsb.grid(column=0, row=1, sticky='ew')


        for col in self.debtmang_tree_header:
                self.debtmang_tree.heading(col, text=col.title(),
                    command=lambda c=col: sortby(tree, c, 0))
                    # adjust the column's width to the header string
                self.debtmang_tree.column(col,
                    width=tkFont.Font().measure(col.title()))

        for item in names_debts:
                    self.debtmang_tree.insert('', 'end', values=item)
                    # adjust column's width if necessary to fit each value
                    for ix, val in enumerate(item):
                        col_w = tkFont.Font().measure(val)
                        if self.debtmang_tree.column(self.debtmang_tree_header[ix],width=None)<col_w:
                            self.debtmang_tree.column(self.debtmang_tree_header[ix], width=col_w)


        #------------------Buttons START-------------------------------
        ttk.Label(debtmang_fr1, text="Info",style='DMMain.TLabel').pack()
        ttk.Label(debtmang_fr1, textvariable=self.DM_infotext,style="DMLbox.TLabel").pack()
        ttk.Label(debtmang_fr1, textvariable=self.DM_infotext2, style="DMLbox.TLabel").pack()
        #empty label
        ttk.Label(debtmang_fr3, text="",style='Empty.TLabel').pack()
        #Other buttons
        debtmang_bnt1 = ttk.Button(debtmang_fr3, text="Poista kaikki velat", style='DMBasic.TButton', command=self.eraseall).pack()
        ttk.Label(debtmang_fr3, text="",style='Empty.TLabel').pack()
        debtmang_bnt2 = ttk.Button(debtmang_fr3, text="Poista valitut velat", style='DMBasic.TButton', command=self.erasesel).pack()
        ttk.Label(debtmang_fr3, text="",style='Empty.TLabel').pack()
        #debtmang_bnt3 = ttk.Button(debtmang_fr3, text="Tarkastele historia tietoja", style='DMBasic.TButton', command=self.paymenthistory).pack()
        #empty label
        ttk.Label(debtmang_fr3, text="",style='Empty.TLabel').pack()
        #Quit button
        debtmang_btn5 = ttk.Button(debtmang_fr3, text='Takaisin', style='DMBasic.TButton', command=debtmang_gui.destroy).pack()
        ttk.Label(debtmang_fr3, text="",style='Empty.TLabel').pack()

    def eraseall(self):
        try:
            top = Toplevel()
            top.geometry("10x10")
            response = messagebox.askokcancel("Varoitus!", "Poistetaanko varmasti kaikki velat?")
            if response == True :

                # Writes payment into separate backup DB --- this is solely for backup purposes only!!!!
                datetime = strftime("%Y-%m-%d %H:%M:%S")
                payname = 'Kaikki velat poistettu'
                erase_all_debt = 0
                conn = sqlite3.connect('history\\' + self.debt_history)
                c = conn.cursor()
                logging.debug('*eraseall* Connecting to DB debt history')
                self.DM_infotext.set("Connected to DB")
                sqlite_insert_with_param = """INSERT INTO debts
                                                        (datetime, name, debt)
                                                        VALUES (?, ?, ?);"""
                data_tuple = (datetime, payname, erase_all_debt)

                # insert row
                c.execute(sqlite_insert_with_param, data_tuple)

                self.DM_infotext.set("Writing to DB done, closing")
                logging.debug('*eraseall* Closing DB' + self.debt_history)
                # commit the changes to db
                conn.commit()
                # close the connection
                conn.close()
                # END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!

                #Deletes all debts from people DB
                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()

                #print("Connected to SQLite DB people")
                self.DM_infotext.set("Connected to DB")

                c.execute('UPDATE people SET debt=0')
                logging.info('*debtmang-eraseall* all debts erased from DB people')
                self.DM_infotext.set("Kaikki velat poistettu!")

                #commit the changes to db
                conn.commit()
                #close the connection
                conn.close()

                #Opens a new DB connection and refreshes the listbox
                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()

                c.execute('SELECT name, debt FROM people ORDER BY name ASC;')
                names_debts = c.fetchall()
                conn.commit()
                #close the connection
                conn.close()

                #selected_item = tree.selection(item) ## get selected item
                for i in self.debtmang_tree.get_children():
                    self.debtmang_tree.delete(i)

                for item in names_debts:
                    self.debtmang_tree.insert('', 'end', values=item)
                    # adjust column's width if necessary to fit each value
                    for ix, val in enumerate(item):
                        col_w = tkFont.Font().measure(val)
                        if self.debtmang_tree.column(self.debtmang_tree_header[ix],width=None)<col_w:
                                self.debtmang_tree.column(self.debtmang_tree_header[ix], width=col_w)
                logging.info('*eraseall* All debts erased')
                top.after(1000, lambda: top.destroy())  # Destroy the widget after 30 seconds
            else:
                top.destroy()

        except Exception as e:
            logging.error('*eraseall* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*eraseall* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.DM_infotext.set("-ERROR-")
            self.DM_infotext2.set("")

    def erasesel(self):

        try:
            items_list = []
            for item in self.debtmang_tree.selection():
                tree_data = self.debtmang_tree.item(item)
                items_list.append(tree_data.get('values', None))

            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            #print("Connected to SQLite DB people")
            self.DM_infotext.set("Connected to DB")

            for item in items_list:
                string = ''
                string += str(item)
                string = "".join(filter(lambda x: not x.isdigit(), string))
                string = string.replace('[','')
                string = string.replace(']','')
                string = string.replace(',','')
                string = string.replace('\'','')
                string = string.replace('.','')
                string = string[:-1]
                string = string.strip()
                #print(string)
                c.execute('UPDATE people SET debt=0.00 WHERE name=?', (string,))
                logging.info('*debtmang-erasesel* selected debts erased from')
                logging.info(string)

            self.DM_infotext.set("Valitut velat poistettu!")
            conn.commit()
            #close the connection
            conn.close()
    #----WRITE ERASED DEBTS INTO HISTORY AS WELL---------------------
            # Writes payment into separate backup DB --- this is solely for backup purposes only!!!!
            datetime = strftime("%Y-%m-%d %H:%M:%S")
            payname = str(items_list)
            erase_all_debt = 0
            conn = sqlite3.connect('history\\' + self.debt_history)
            c = conn.cursor()
            logging.debug('*eraseall* Connecting to DB debt history')
            self.DM_infotext.set("Connected to DB")
            write_text = "Valitut velat poistettu:"
            sqlite_insert_with_param = """INSERT INTO debts
                                                                                (datetime, name, debt)
                                                                                VALUES (?, ?, ?);"""
            data_tuple = (datetime, write_text, erase_all_debt)

            # insert row
            c.execute(sqlite_insert_with_param, data_tuple)

            sqlite_insert_with_param = """INSERT INTO debts
                                                                    (datetime, name, debt)
                                                                    VALUES (?, ?, ?);"""
            data_tuple = (datetime, payname, erase_all_debt)

            # insert row
            c.execute(sqlite_insert_with_param, data_tuple)

            logging.debug('*erasesel* Closing DB' + self.debt_history)
            # commit the changes to db
            conn.commit()
            # close the connection
            conn.close()
            # END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!

            #Opens a new DB connection and refreshes the listbox
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            c.execute('SELECT name, debt FROM people ORDER BY name ASC;')
            names_debts = c.fetchall()
            conn.commit()
            #close the connection
            conn.close()
            self.DM_infotext.set("Valitut velat poistettu")
            #selected_item = tree.selection(item) ## get selected item
            for i in self.debtmang_tree.get_children():
                self.debtmang_tree.delete(i)

            for item in names_debts:
                    self.debtmang_tree.insert('', 'end', values=item)
                    # adjust column's width if necessary to fit each value
                    for ix, val in enumerate(item):
                        col_w = tkFont.Font().measure(val)
                        if self.debtmang_tree.column(self.debtmang_tree_header[ix],width=None)<col_w:
                            self.debtmang_tree.column(self.debtmang_tree_header[ix], width=col_w)

        except Exception as e:
            logging.error('*erasesel* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*erasesel* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.DM_infotext.set("-ERROR-")
            self.DM_infotext2.set("")

    def paymenthistory(self):
        os.system('python history.py')

    def exporttoexl(self):
        try:
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            #print("Connected to SQLite DB people")
            self.infotext.set("Connected to DB")
            #create table

            c.execute('SELECT name, debt FROM people WHERE debt > 0 ORDER BY name ASC;')
            names_debts_xls = c.fetchall()

            self.infotext.set("DB read, closing DB")
            conn.commit()

            #close the connection
            conn.close()
            savedate = strftime("%d-%m-%Y")
            filename = '{}_velkalista.csv'
            filename = filename.format(savedate)
            username = os.getlogin()    # Fetch username
            #file = open(f'C:\\Users\\{username}\\Desktop\\PC info.txt','w')
            f = open(f'C:\\Users\\{username}\\Desktop\\' + filename,'w')
            #f = open(filename, 'w')
            header_string = 'Name, debt\n'
            f.write(header_string)
            for item in names_debts_xls:
                string = ''
                string += str(item)
                string = string.replace('[','')
                string = string.replace(']','')
                string = string.replace('(','')
                string = string.replace(')','')
                string = string.replace('\'','')
                string += '\n'
                #print(string)
                f.write(string)

            f.close()
            self.infotext.set("Excel lista luotu työpöydälle\n" + filename)

        except Exception as e:
            logging.error('*exporttoexl* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*exporttoexl* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.infotext.set("-ERROR-")

    def editpayment(self, arg):
        windowWidth = gui.winfo_reqwidth()
        windowHeight = gui.winfo_reqheight()
        # print("Width", windowWidth, "Height", windowHeight)
        positionRight = int(gui.winfo_screenwidth() / 2 - windowWidth / 2)
        positionDown = int(gui.winfo_screenheight() / 2 - windowHeight / 2)
        self.editpayment_top = Toplevel()
        self.editpayment_top.geometry("+{}+{}".format(positionRight, positionDown))
        self.editpayment_top.attributes("-topmost", True)
        editpayment_fr = ttk.Frame(self.editpayment_top, relief="raised", borderwidth=1)
        editpayment_fr.pack(expand=0, fill="both")
        items_list = []
        for item in self.debtmang_tree.selection():
            self.debtmang_tree_data = self.debtmang_tree.item(item)
            items_list.append(self.debtmang_tree_data.get('values', None))

            for item in items_list:
                string = ''
                string += str(item)
                string = "".join(filter(lambda x: not x.isdigit(), string))
                string = string.replace('[','')
                string = string.replace(']','')
                string = string.replace(',','')
                string = string.replace('\'','')
                string = string.replace('.','')
                self.string = string.strip()

                self.EP_infotext.set(self.string)

        self.EP_infotext2.set("")
        ep_lbl1 = ttk.Label(editpayment_fr, text='Lisää maksu:',style="DMLbox.TLabel").pack()
        ep_lbl2 = ttk.Label(editpayment_fr, textvariable=self.EP_infotext,style="DMLbox.TLabel").pack()
        ep_lbl3 = ttk.Label(editpayment_fr, textvariable=self.EP_infotext2,style="DMLbox.TLabel").pack()
        self.ep_ent = ttk.Entry(editpayment_fr, justify = 'center')
        self.ep_ent.pack()
        self.ep_btn1 = ttk.Button(editpayment_fr, text="Tallenna", style='DMBasic.TButton', command=self.epsave)
        self.ep_btn1.pack()
        self.ep_btn2 = ttk.Button(editpayment_fr, text="Takaisin", style='DMBasic.TButton', command=self.editpayment_top.destroy)
        self.ep_btn2.pack()

    def epsave(self):
        try:
            check_input = self.ep_ent.get()

            if check_input == "":
                self.EP_infotext2.set('Syötä numero!\n Esim. 2.3')
            else:
                #self.ep_true = True
                input = self.ep_ent.get()
                input_round = float(input)
                input_round = round(input_round,2)
                #print(input_round)
                payname = self.string
                #Writes payment into people DB
                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()
                #print("Connected to SQLite DB people")
                self.EP_infotext.set("Connected to DB")

                c.execute('UPDATE people SET payment=? WHERE name=?', (input, payname))
                input_str = str(input_round)
                self.EP_infotext.set("Maksu " + input_str + " kirjattu")
                #commit the changes to db
                conn.commit()
                #close the connection
                conn.close()
#--------------WRITES PAYMENT TO HISTORY DB--------------------------------------
                date_now = strftime("%m-%Y")
                date_now = date_now + "_debts.db"
                paytime = strftime('%H:%M:%S')
                paydate = strftime("%d-%m-%Y")
                datetime = strftime("%Y-%m-%d %H:%M:%S")
                payment_to_save = input_round * -1
                #print(payment_to_save)
                if date_now != self.debt_history:
                    logging.info('Month change detected, creating a new history DB')
                    self.debt_history = date_now + "_debts.db"
                    logging.info(self.debt_history)
                    conn = sqlite3.connect('history\\' + self.debt_history)
                    c = conn.cursor()
                    # create table 'people'
                    c.execute('''CREATE TABLE IF NOT EXISTS debts
                                                                         (datetime, name, debt, payment)''')

                    conn.commit()
                    # close the connection
                    conn.close()
                else:

                    # Writes payment into separate backup DB --- this is solely for backup purposes only!!!!
                    conn = sqlite3.connect('history\\' + self.debt_history)
                    c = conn.cursor()
                    logging.debug('*process* Connecting to DB debt history')
                    self.DM_infotext.set("Connected to DB")
                    sqlite_insert_with_param = """INSERT INTO debts
                                                                    (datetime, name, debt)
                                                                    VALUES (?, ?, ?);"""
                    data_tuple = (datetime, payname, payment_to_save)

                    # insert row
                    c.execute(sqlite_insert_with_param, data_tuple)

                    self.DM_infotext.set("Writing to DB done, closing")
                    logging.debug('*epsave* Closing DB' + self.debt_history)
                    # commit the changes to db
                    conn.commit()
                    # close the connection
                    conn.close()
#----------------------- END OF --- Writes payment into separate debts DB --- this is solely for backup purposes only!!!!

                conn = sqlite3.connect('DB\people.db')
                c = conn.cursor()
                #print("Connected to SQLite DB people")
                #infotext.set("Connected to DB")

                c.execute('SELECT name, debt FROM people ORDER BY name ASC;')
                names_debts = c.fetchall()

                #print('DB fetchall:')
                #print(names_debts)

                #print("DB read, closing DB")
                #infotext.set("DB read, closing DB")
                conn.commit()
                #close the connection
                conn.close()

                for i in self.debtmang_tree.get_children():
                    self.debtmang_tree.delete(i)

                #print('Listbox names deleted')

                for item in names_debts:
                    self.debtmang_tree.insert('', 'end', values=item)
                    # adjust column's width if necessary to fit each value
                    for ix, val in enumerate(item):
                        col_w = tkFont.Font().measure(val)
                        if self.debtmang_tree.column(self.debtmang_tree_header[ix],width=None)<col_w:
                                self.debtmang_tree.column(self.debtmang_tree_header[ix], width=col_w)

            self.editpayment_top.destroy()
            self.process()

        except Exception as e:
            logging.error('*epsave* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*epsave* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.DM_infotext.set("-ERROR-")
            self.DM_infotext2.set("")

    def process(self):
    #SAVE TO HISTORY DB IS DONE IN DEF EPSAVE
        try:

            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            #print("Connected to SQLite DB people")
            self.DM_infotext.set("Connected to DB")

            c.execute('SELECT name, debt - payment FROM people ORDER BY name ASC;')
            names_debts = c.fetchall()

            #print('DB fetchall:')
            #print(names_debts)
            name_str = ''
            debt_str = '0.00'
            for row in names_debts:
                name_str = str(row[0])
                debt_str = float(row[1])
                debt_str = round(debt_str,2)
                #print(name_str)
                #print(debt_str)
                c.execute('UPDATE people SET debt=? WHERE name=?', (debt_str, name_str,))
                c.execute('UPDATE people SET payment=0 WHERE name=?', (name_str,))
            #print("DB read, closing DB")
            self.DM_infotext.set("DB read, closing DB")
            conn.commit()
            #close the connection
            conn.close()

            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            #print("Connected to SQLite DB people")
            self.DM_infotext.set("Updating list")

            c.execute('SELECT name, debt FROM people ORDER BY name ASC;')
            names_debts = c.fetchall()

            #print('DB fetchall:')
            #print(names_debts)

            #print("DB read, closing DB")
            #infotext.set("DB read, closing DB")
            conn.commit()
            #close the connection
            conn.close()

            for i in self.debtmang_tree.get_children():
                self.debtmang_tree.delete(i)

            #print('Listbox names deleted')

            for item in names_debts:
                self.debtmang_tree.insert('', 'end', values=item)
                # adjust column's width if necessary to fit each value
                for ix, val in enumerate(item):
                    col_w = tkFont.Font().measure(val)
                    if self.debtmang_tree.column(self.debtmang_tree_header[ix],width=None)<col_w:
                            self.debtmang_tree.column(self.debtmang_tree_header[ix], width=col_w)


            self.DM_infotext.set(self.string)
            self.DM_infotext2.set("Maksu tallennettu.")
        except Exception as e:
            logging.error('*epsave-process* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*epsave-process* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.DM_infotext.set("-ERROR-")

#-------------- END OF DEBT MANAGEMENT ----------------------------------------
    def print(self):
        try:
            conn = sqlite3.connect('DB\people.db')
            c = conn.cursor()

            logging.debug('*print* Connected to DB people')
            self.infotext.set("Connected to DB")
            #create table

            c.execute('SELECT name, debt FROM people WHERE debt > 0 ORDER BY name ASC;')
            names_debts_toprint = c.fetchall()

            logging.debug('*print* DB fetchall')
            logging.debug(names_debts_toprint)

            #print("DB read, closing DB")
            self.infotext.set("DB read, closing DB")
            logging.debug('*print* DB read, closing DB data\people.db')
            conn.commit()

            #close the connection
            conn.close()
            logging.debug('*print* Creating docx document')
            logging.info('*print* Creating docx document')
            document = Document()

            document.add_picture('data\print_logo.png')

            document.add_heading('IOK Kahvikerhon velkalista', 0)

            p = document.add_paragraph('Tili ja maksutiedot yms. tähän ')
            p.add_run('bold').bold = True
            p.add_run(' and some ')
            p.add_run('italic.').italic = True
            month_year = strftime("%d-%m-%Y")
            document.add_heading('Kahvivelat ' + month_year, level=1)

            records = (names_debts_toprint)

            table = document.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Nimi'
            hdr_cells[1].text = 'Velka (€)'
            hdr_cells[2].text = 'Kuittaus'

            for name, debt in records:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = str(debt)
                #row_cells[2].text = desc

            document.add_page_break()


            username = os.getlogin()    # Fetch username
            #file = open(f'C:\\Users\\{username}\\Desktop\\PC info.txt','w')

            document.save(f'C:\\Users\\{username}\\Desktop\\'+ month_year + '_velkalista.docx')
            logging.debug('*print* Docx document saved')
            self.infotext.set("Word asiakirja luotu työpöydälle:\n" + month_year + '_velkalista.docx')
            logging.info('*print* Docx document saved')

        except Exception as e:
            logging.error('*print* ERROR. Reason: %s' % (e))
            username = os.getlogin()  # Fetch username
            date = strftime("%d-%m-%Y")
            print('Creating error log')
            error_to_write = '*print* -ERROR- Reason: %s' % (e)
            f = open(f"C:\\Users\\{username}\\Desktop\\" + date + "_error_info.txt", "w")
            f.write(ver)
            f.write('\n')
            f.write(error_to_write)
            f.write('\n')
            f.close()
            error_info = str(date + '_error_info.txt')
            messagebox.showerror("Error",
                                 "Something went wrong - description (" + error_info + ") created to desktop")
            self.infotext.set("-ERROR-")


    def __init__(self):
        print('Starting program')

        #Variables
        self.infotext = StringVar("")
        self.infotext2 = StringVar("")
        self.infotext.set("")
        self.debugvar = IntVar()
        self.Paybkupvar = IntVar()
        self.debt_to_add = 0
        #Check that needed paths exists
        prog_path = os.path.dirname(os.path.abspath(__file__ ))
        #print(prog_path)
        print('Checking prerequisites')
        settings_path = prog_path+'\data'
        setup_file = prog_path+'\data\settings.file'
        DB_path = prog_path+'\DB'
        history_path = prog_path+'\history'
        log_path = prog_path+'\log'

        is_settings_path = os.path.isdir(settings_path)
        #print(is_settings_path)
        if is_settings_path == False:
            os.mkdir(settings_path)
            print('No data folder found, creating a folder')

        is_settings_File = os.path.isfile(setup_file)
        #print(is_settings_File)
        if is_settings_File == False:
            print('No settings file found, creating one with default values')
            debugstg = '0'
            backupstg = '1'
            debugtow = 'Debug:'+debugstg
            backuptow = 'Backup:'+backupstg

            f = open("data\settings.file", "w")
            f.write(debugtow)
            f.write('\n')
            f.write(backuptow)
            f.write('\n')
            f.close()

        is_DB_dir = os.path.isdir(DB_path)
        #print(is_DB_dir)
        if is_DB_dir == False:
            os.mkdir(DB_path)
            print('No DB folder found, creating a folder')

        is_history_path_dir = os.path.isdir(history_path)
        #print(is_history_path_dir)
        if is_history_path_dir == False:
            os.mkdir(history_path)
            print('No history folder found, creating a folder')

        is_log_path_dir = os.path.isdir(log_path)
        if is_log_path_dir == False:
            os.mkdir(log_path)
            print('No log folder found, creating a folder')

        #Open settings
        f = open("data\settings.file", "r")
        debug = f.readline()
        paybkup = f.readline()
        f.close()
        debug = debug[6]
        paybkup = paybkup[7]
        print('Settings read')

        self.curr_mID = subprocess.check_output('wmic csproduct get uuid').decode().split('\n')[1].strip()
        #print(self.curr_mID)

        self.debugvar.set(debug)
        self.Paybkupvar.set(paybkup)

        #---------------------Define Logging level------------------------------------
        #logging.debug('This is a debug message')
        #logging.info('This is an info message')
        #logging.warning('This is a warning message')
        #logging.error('This is an error message')
        #logging.critical('This is a critical message')
        logdate = strftime("%d%m%Y")
        if debug == '1':
            logginglevel = logging.DEBUG
            print('Logging level DEBUG')
        else:
            logginglevel = logging.INFO
            print('Logging level INFO')

        logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',filename= ('log\ '+logdate+'_applog.log'), filemode='w',level=logginglevel)
        logging.info('Settings read')
        logging.debug('Debug:'+debug)
        logging.debug('Pay back up:'+paybkup)
        self.launchergui()
#-------------Information section starts here ----------------------
# Set version number
ver = "PV2.11"
#Create a new instance
#Define theme
gui = ThemedTk(theme="blue") #Tested and working themes: aquablue, black, blue, clam, classic, clearlooks, equilux, gtk2, itft1, keramik, smog,
#Set GUI parameters - Title
gui.title("Kahvikassa " + ver)
#Set GUI parameters - Is window resizable
#gui.resizable(True, True)
#Set GUI parameters - Window size x+y
gui.minsize(width=800, height=600)
gui.maxsize(width=1440, height=828)
#gui.geometry("800x600")

positionRight = int(gui.winfo_screenwidth() / 4)
positionDown = int(gui.winfo_screenheight() / 4)

# Positions the window in the center of the page.
gui.geometry("+{}+{}".format(positionRight, positionDown))
#Set GUI parameters - Background window color
#gui.configure(bg="white")
view = mainprogram()
#Start the GUI/Create main loop
gui.mainloop()
